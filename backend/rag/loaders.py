"""Document loading and text extraction."""

from pathlib import Path

import fitz
from docx import Document as DocxDocument

from backend.rag.errors import (
    DocumentNotFoundError,
    DocumentTooLargeError,
    EmptyDocumentError,
    UnsupportedDocumentError,
)
from backend.rag.models import LoadedDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def load_document(
    document_path: str | Path,
    *,
    max_size_mb: int = 20,
) -> LoadedDocument:
    """Load a supported document and extract normalized text."""

    path = Path(document_path).expanduser().resolve()

    if not path.exists():
        raise DocumentNotFoundError(
            f"Файл не знайдено: {path}"
        )

    if not path.is_file():
        raise DocumentNotFoundError(
            f"Шлях не є файлом: {path}"
        )

    extension = path.suffix.casefold()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))

        raise UnsupportedDocumentError(
            f"Формат '{extension or 'без розширення'}' не підтримується. "
            f"Доступні формати: {supported}."
        )

    size_bytes = path.stat().st_size
    max_size_bytes = max_size_mb * 1024 * 1024

    if size_bytes > max_size_bytes:
        raise DocumentTooLargeError(
            f"Файл завеликий: {size_bytes / 1024 / 1024:.1f} MB. "
            f"Максимальний розмір: {max_size_mb} MB."
        )

    if extension == ".pdf":
        content, page_count = _load_pdf(path)
    else:
        content, page_count = _load_docx(path)

    normalized_content = normalize_text(content)

    if not normalized_content:
        raise EmptyDocumentError(
            f"Не вдалося витягнути текст із документа: {path.name}"
        )

    return LoadedDocument(
        path=path,
        filename=path.name,
        extension=extension,
        content=normalized_content,
        size_bytes=size_bytes,
        page_count=page_count,
    )


def _load_pdf(path: Path) -> tuple[str, int]:
    """Extract text from a PDF document."""

    pages: list[str] = []

    try:
        with fitz.open(path) as document:
            for page in document:
                page_text = page.get_text("text")

                if page_text.strip():
                    pages.append(page_text)

            return "\n\n".join(pages), document.page_count
    except (fitz.FileDataError, RuntimeError, ValueError) as error:
        raise EmptyDocumentError(
            f"Не вдалося прочитати PDF '{path.name}': {error}"
        ) from error


def _load_docx(path: Path) -> tuple[str, int | None]:
    """Extract paragraphs and table values from a DOCX document."""

    try:
        document = DocxDocument(path)
    except Exception as error:
        raise EmptyDocumentError(
            f"Не вдалося прочитати DOCX '{path.name}': {error}"
        ) from error

    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                blocks.append(" | ".join(values))

    return "\n\n".join(blocks), None


def normalize_text(value: str) -> str:
    """Normalize extracted text without destroying paragraph boundaries."""

    normalized_lines: list[str] = []
    previous_line_was_empty = False

    for raw_line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = " ".join(raw_line.split())

        if not line:
            if normalized_lines and not previous_line_was_empty:
                normalized_lines.append("")

            previous_line_was_empty = True
            continue

        normalized_lines.append(line)
        previous_line_was_empty = False

    return "\n".join(normalized_lines).strip()