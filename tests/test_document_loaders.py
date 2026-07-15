from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument

from backend.rag.errors import (
    DocumentNotFoundError,
    UnsupportedDocumentError,
)
from backend.rag.hashing import calculate_file_hash
from backend.rag.loaders import load_document


def test_load_docx_extracts_paragraphs_and_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "example.docx"

    document = DocxDocument()
    document.add_heading("Project report", level=1)
    document.add_paragraph("First paragraph.")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Kuro"

    document.save(path)

    loaded = load_document(path)

    assert loaded.filename == "example.docx"
    assert loaded.extension == ".docx"
    assert "Project report" in loaded.content
    assert "First paragraph." in loaded.content
    assert "Name | Kuro" in loaded.content


def test_load_pdf_extracts_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "example.pdf"

    document = fitz.open()
    page = document.new_page()
    page.insert_text(
        (72, 72),
        "Kuro PDF document",
    )
    document.save(path)
    document.close()

    loaded = load_document(path)

    assert loaded.filename == "example.pdf"
    assert loaded.extension == ".pdf"
    assert loaded.page_count == 1
    assert "Kuro PDF document" in loaded.content


def test_missing_document_raises_error(
    tmp_path: Path,
) -> None:
    with pytest.raises(DocumentNotFoundError):
        load_document(tmp_path / "missing.pdf")


def test_unsupported_extension_raises_error(
    tmp_path: Path,
) -> None:
    path = tmp_path / "example.txt"
    path.write_text("Hello", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentError):
        load_document(path)


def test_file_hash_is_stable(
    tmp_path: Path,
) -> None:
    path = tmp_path / "document.pdf"
    path.write_bytes(b"same content")

    first_hash = calculate_file_hash(path)
    second_hash = calculate_file_hash(path)

    assert first_hash == second_hash
    assert len(first_hash) == 64